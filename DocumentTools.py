from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Mm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PAGE_SIZE = {
    "A4": (Mm(210), Mm(297)),
    "A5": (Mm(148), Mm(210)),
}

def set_margin(section, margin):
    section.left_margin, section.right_margin, section.top_margin, section.bottom_margin = [margin]*4

def create_document():
    return Document()

def configure_parallel(doc, page_size="A4", portrait=True, margin=Mm(15)):
    """ Two column page """
    section = doc.sections[-1]
    section.page_width, section.page_height = PAGE_SIZE[page_size if portrait else page_size[::-1]]
    set_margin(section, margin)
    table = doc.add_table(rows=1, cols=2)
    return table.cell(0,0), table.cell(0,1)

def configure_singular(doc, page_size="A4", portrait=True, margin=Mm(15)):
    """ One column page """
    section = doc.sections[-1]
    section.page_width, section.page_height = PAGE_SIZE[page_size if portrait else page_size[::-1]]
    set_margin(section, margin)
    table = doc.add_table(rows=1, cols=1)
    return table.cell(0,0)

def add_styles(document):
    note = document.styles.add_style("note", style_type = WD_STYLE_TYPE.CHARACTER)
    note.font.superscript = True

    label = document.styles.add_style("label", style_type = WD_STYLE_TYPE.CHARACTER)
    label.font.superscript = True

    small_caps = document.styles.add_style("sc", style_type = WD_STYLE_TYPE.CHARACTER)
    small_caps.font.small_caps = True
    small_caps_german = document.styles.add_style("nd", style_type = WD_STYLE_TYPE.CHARACTER)
    small_caps_german.font.small_caps = True

    words = document.styles.add_style("w", style_type = WD_STYLE_TYPE.CHARACTER) # Used in lversions like 2134 (Dhivehi)

    paragraph_caps = document.styles.add_style("pc", style_type = WD_STYLE_TYPE.PARAGRAPH)
    paragraph_caps.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_caps.font.small_caps = True

    words_of_Jesus = document.styles.add_style("wj", style_type = WD_STYLE_TYPE.CHARACTER)
    words_of_Jesus.font.color.rgb = RGBColor(255,0,0)

    # Regular content
    document.styles.add_style("content", style_type = WD_STYLE_TYPE.CHARACTER)
    document.styles.add_style("pn", style_type = WD_STYLE_TYPE.CHARACTER)

    heading = document.styles.add_style("heading", style_type = WD_STYLE_TYPE.CHARACTER)
    heading.font.bold = True

    chapter_heading = document.styles.add_style("chapter_heading", style_type = WD_STYLE_TYPE.CHARACTER)
    chapter_heading.font.bold = True
    chapter_heading.font.size = Pt(16)

    # Regular paragraphs
    document.styles.add_style("p", style_type = WD_STYLE_TYPE.PARAGRAPH)
    document.styles.add_style("m", style_type = WD_STYLE_TYPE.PARAGRAPH)

    quote_level_1 = document.styles.add_style("q1", style_type = WD_STYLE_TYPE.PARAGRAPH)
    quote_level_1.paragraph_format.left_indent = Pt(15)
    quote_level_1.paragraph_format.space_after = Pt(0)

    quote_level_2 = document.styles.add_style("q2", style_type = WD_STYLE_TYPE.PARAGRAPH)
    quote_level_2.paragraph_format.left_indent = Pt(30)
    quote_level_2.paragraph_format.space_after = Pt(0)

    quote_level_3 = document.styles.add_style("q3", style_type = WD_STYLE_TYPE.PARAGRAPH)
    quote_level_3.paragraph_format.left_indent = Pt(40)
    quote_level_3.paragraph_format.space_after = Pt(0)

    # Annotations, as in Psalm 119, and section headings
    quote_annotation = document.styles.add_style("qa", style_type = WD_STYLE_TYPE.PARAGRAPH)
    quote_annotation.paragraph_format.space_before = Pt(12)
    quote_annotation.paragraph_format.space_after = Pt(0)

    blank_line = document.styles.add_style("blank_line", style_type = WD_STYLE_TYPE.PARAGRAPH)

    table_vertical_two_column = document.styles.add_style("table_vertical_two_column", style_type = WD_STYLE_TYPE.TABLE)
    table_vertical_two_column

    footnotes = document.styles.add_style("footnotes", style_type = WD_STYLE_TYPE.PARAGRAPH)
    footnotes.font.size = Pt(9)

    copyright = document.styles.add_style("copyright", style_type = WD_STYLE_TYPE.PARAGRAPH)
    copyright.font.size = Pt(9)
    copyright.font.italic = True


def add_end_break(doc, break_type=WD_BREAK.COLUMN):
    runs = doc.paragraphs[-1].runs
    if len(runs) == 0:
        doc.paragraphs[-1].add_run()
    doc.paragraphs[-1].runs[-1].add_break(break_type)