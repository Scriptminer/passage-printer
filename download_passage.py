from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Mm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re
import requests
from bs4 import BeautifulSoup

CACHE_DIR = "cache"
URL_PATTERN = "https://www.bible.com/bible/{version_id}/{book_code}.{chapter}"
PAGE_SIZE = {
    "A4": (Mm(210), Mm(297)),
    "A5": (Mm(148), Mm(210)),
}

def set_margin(section, margin):
    section.left_margin, section.right_margin, section.top_margin, section.bottom_margin = [margin]*4

def configure_parallel(doc, page_size="A4", portrait=True, margin=Mm(15)):
    """ Two column page """
    section = doc.sections[-1]
    section.page_width, section.page_height = PAGE_SIZE[page_size if portrait else page_size[::-1]]
    set_margin(section, margin)
    table = doc.add_table(rows=1, cols=2)
    return table.cell(0,0), table.cell(0,1)

def configure_singular(doc, page_size="A4", portrait=True, margin=Mm(15)):
    """ One column page """
    section = doc.sections[-1]
    section.page_width, section.page_height = PAGE_SIZE[page_size if portrait else page_size[::-1]]
    set_margin(section, margin)
    table = doc.add_table(rows=1, cols=1)
    return table.cell(0,0)

def add_styles(document):
    note = document.styles.add_style("note", style_type = WD_STYLE_TYPE.CHARACTER)
    note.font.superscript = True

    label = document.styles.add_style("label", style_type = WD_STYLE_TYPE.CHARACTER)
    label.font.superscript = True

    small_caps = document.styles.add_style("sc", style_type = WD_STYLE_TYPE.CHARACTER)
    small_caps.font.small_caps = True
    small_caps_german = document.styles.add_style("nd", style_type = WD_STYLE_TYPE.CHARACTER)
    small_caps_german.font.small_caps = True

    words = document.styles.add_style("w", style_type = WD_STYLE_TYPE.CHARACTER) # Used in lversions like 2134 (Dhivehi)

    paragraph_caps = document.styles.add_style("pc", style_type = WD_STYLE_TYPE.PARAGRAPH)
    paragraph_caps.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_caps.font.small_caps = True

    words_of_Jesus = document.styles.add_style("wj", style_type = WD_STYLE_TYPE.CHARACTER)
    words_of_Jesus.font.color.rgb = RGBColor(255,0,0)

    # Regular content
    document.styles.add_style("content", style_type = WD_STYLE_TYPE.CHARACTER)
    document.styles.add_style("pn", style_type = WD_STYLE_TYPE.CHARACTER)

    heading = document.styles.add_style("heading", style_type = WD_STYLE_TYPE.CHARACTER)
    heading.font.bold = True

    chapter_heading = document.styles.add_style("chapter_heading", style_type = WD_STYLE_TYPE.CHARACTER)
    chapter_heading.font.bold = True
    chapter_heading.font.size = Pt(16)

    # Regular paragraphs
    document.styles.add_style("p", style_type = WD_STYLE_TYPE.PARAGRAPH)
    document.styles.add_style("m", style_type = WD_STYLE_TYPE.PARAGRAPH)

    quote_level_1 = document.styles.add_style("q1", style_type = WD_STYLE_TYPE.PARAGRAPH)
    quote_level_1.paragraph_format.left_indent = Pt(15)
    quote_level_1.paragraph_format.space_after = Pt(0)

    quote_level_2 = document.styles.add_style("q2", style_type = WD_STYLE_TYPE.PARAGRAPH)
    quote_level_2.paragraph_format.left_indent = Pt(30)
    quote_level_2.paragraph_format.space_after = Pt(0)

    quote_level_3 = document.styles.add_style("q3", style_type = WD_STYLE_TYPE.PARAGRAPH)
    quote_level_3.paragraph_format.left_indent = Pt(40)
    quote_level_3.paragraph_format.space_after = Pt(0)

    # Annotations, as in Psalm 119, and section headings
    quote_annotation = document.styles.add_style("qa", style_type = WD_STYLE_TYPE.PARAGRAPH)
    quote_annotation.paragraph_format.space_before = Pt(12)
    quote_annotation.paragraph_format.space_after = Pt(0)

    blank_line = document.styles.add_style("blank_line", style_type = WD_STYLE_TYPE.PARAGRAPH)

    table_vertical_two_column = document.styles.add_style("table_vertical_two_column", style_type = WD_STYLE_TYPE.TABLE)
    table_vertical_two_column

    footnotes = document.styles.add_style("footnotes", style_type = WD_STYLE_TYPE.PARAGRAPH)
    footnotes.font.size = Pt(9)

    copyright = document.styles.add_style("copyright", style_type = WD_STYLE_TYPE.PARAGRAPH)
    copyright.font.size = Pt(9)
    copyright.font.italic = True

def resolve_version(version_name):
    if type(version_name) == int:
        return version_name
    language_info = [None, "", None]
    with open("custom_version_info.txt") as f:
        for line in f.readlines():
            parts = line.strip().split(";")
            if version_name.lower() in parts[1].lower():
                language_info = parts
    return language_info[0]

def add_end_break(doc, break_type=WD_BREAK.COLUMN):
    runs = doc.paragraphs[-1].runs
    if len(runs) == 0:
        doc.paragraphs[-1].add_run()
    doc.paragraphs[-1].runs[-1].add_break(break_type)

def format_url(version_id, book_code,chapter):
    return URL_PATTERN.format(version_id=version_id,book_code=book_code,chapter=chapter)

class FootnoteHandler:
    def __init__(self):
        self.text_notes = []
        self.next_note_label = "a"
        pass
    
    def add_note(self, footnote_section):
        footnote = footnote_section.getText()[1:] # Need to extract italics etc in future
        label = self.next_note_label
        self.text_notes.append( (label, footnote) )
        self.update_next_note_label()
        return label
    
    def update_next_note_label(self):
        self.next_note_label = chr(ord(self.next_note_label)+1)
    
    def print_notes(self):
        out = "\n".join([f"{label}: {note}" for label, note in self.text_notes])
        return out

class PassagePointer:
    BEFORE_START = 0
    IN_PASSAGE = 1
    AFTER_END = 2

    def __init__(self, start_verse, end_verse):
        self.start_verse = start_verse
        self.current_verse = start_verse
        self.end_verse = end_verse
        self.last_section = None
        self.state = PassagePointer.BEFORE_START
    
    def update_last_section(self, section):
        self.last_section = section
    
    def update_state(self, verse_text):
        self.set_current_verse(verse_text)
        if self.current_verse >= self.start_verse:
            self.state = PassagePointer.IN_PASSAGE
        if self.end_verse != -1 and self.current_verse > self.end_verse:
            self.state = PassagePointer.AFTER_END

    def set_current_verse(self, verse_text):
        self.current_verse = int(verse_text)

    def __str__(self):
        return str(self.current_verse)
    
def get_passage(version_id, book_code, chapter):
    fname = f"{version_id}.{book_code}.{chapter}.html"
    yvReader = None
    if fname in os.listdir(CACHE_DIR):
        print(f"Loading {fname} from cache.")
        with open(os.path.join(CACHE_DIR, fname), "r") as f:
            yvReader = BeautifulSoup(f.read())
    else:
        url = format_url(version_id, book_code, chapter)
        print(f"Loading from {url}")
        response = requests.get(url)
        page = response.content
        soup = BeautifulSoup(page)
        yvReader = soup.find("div", class_=re.compile("ChapterContent_yv-bible-text"))

        # Save to cache
        with open(os.path.join(CACHE_DIR, fname), "w", encoding="utf-8") as f:
            f.write(str(yvReader))
    
    return yvReader

class Version:
    version_id = None
    name = "Holy Bible"
    custom_copyright_statement = None

    def __init__(self, version_id, version_information_file="custom_version_info.txt"):
        self.version_id = version_id
        with open(version_information_file) as f:
            for line_no, line in enumerate(f.readlines()):
                parts = line.strip().split(";")
                if len(parts) != 3:
                    print(f"Invalid version info file at line {line_no}. Expected '<version_id>;<name>;<copyright statement>', got '{line.strip()}'.")
                    continue
                if parts[0] == str(version_id):
                    self.name, self.custom_copyright_statement = parts[1:]
        if self.custom_copyright_statement == None:
            print(f"Version information for version_id {self.version_id} missing, using defaults.")
    
    def __str__(self):
        return str(self.version_id)

class Passage:
    def __init__(self, version_id="113", book_code="MRK", chapter="1", start_verse=1, end_verse=-1):
        self.version = Version(version_id)
        self.book_code, self.chapter, self.start_verse, self.end_verse = book_code, chapter, start_verse, end_verse
        print(f"Loading passage: {self}...")

        yvReader = get_passage(version_id, book_code, chapter)
        self.copyright_statement = self.get_copyright_statement(yvReader)
        self.readmore_statement = f"Read more at {format_url(self.version.version_id, self.book_code, self.chapter)}"
        self.footnotes_handler = FootnoteHandler()
        self.passage_pointer = PassagePointer(start_verse, end_verse)

        self.chapter_content = yvReader.find("div", class_=re.compile("ChapterContent_chapter"))
        self.chapter_title = yvReader.find("div", class_=re.compile("ChapterContent_reader")).find("h1").getText()

        print(f"Loading passage {self} complete.")
    
    def write_to(self, doc, config=None):
        doc.paragraphs[0].add_run(self.version.name, style="heading")
        doc.add_paragraph().add_run(self.chapter_title, style="chapter_heading")

        for chapter_section in self.chapter_content.find_all(recursive=False):
            section_type = re.findall("^ChapterContent_([a-zA-Z0-9]*)_*.*$", chapter_section["class"][0])[0]
            if section_type in ["p", "q1", "q2", "q3", "qa", "d", "pc", "m"]:
                doc_paragraph = {"paragraph": False, "doc": doc, "create": lambda doc: doc.add_paragraph(style=section_type)}
                if self.passage_pointer.state == PassagePointer.IN_PASSAGE:
                    doc_paragraph["paragraph"] = doc_paragraph["create"](doc)
                
                for paragraph_section in chapter_section.find_all(recursive=False):
                    paragraph_section_type = re.findall("^ChapterContent_([a-zA-Z0-9]*)_*.*$", paragraph_section["class"][0])[0]
                    if paragraph_section_type == "verse":
                        for verse_section in paragraph_section.find_all(recursive=False):
                            self.add_verse_section(doc_paragraph, verse_section, self.footnotes_handler, self.passage_pointer)
                            if self.passage_pointer.state == PassagePointer.AFTER_END: break
                    elif paragraph_section_type in ["content", "heading"]:
                        self.add_verse_section(doc_paragraph, paragraph_section, self.footnotes_handler, self.passage_pointer)
                    else:
                        print(f"Unexpected part of section '{paragraph_section}' found.")
                    if self.passage_pointer.state == PassagePointer.AFTER_END: break
                if self.passage_pointer.state == PassagePointer.AFTER_END: break
            elif section_type in ["s", "s1"]:
                self.add_heading_section(doc, chapter_section.getText(), self.passage_pointer)
                
            elif section_type == "b":
                if self.passage_pointer.IN_PASSAGE:
                    doc.add_paragraph(style="blank_line")

            elif section_type == "label":
                pass # Skip labels

            else:
                print(f"Unexpected section type '{section_type}' found. Section contents: {chapter_section}")

        doc.add_paragraph(self.footnotes_handler.print_notes(), style="footnotes")
        doc.add_paragraph(self.copyright_statement + "\n" + self.readmore_statement, style="copyright")

    def add_verse_section(self, paragraph, verse_section_data, footnotes_handler, passage_pointer):
        extracted_section_type = re.findall("^ChapterContent_([a-zA-Z0-9]*)_*.*$", verse_section_data["class"][0])
        verse_section_type = extracted_section_type[0] if len(extracted_section_type) > 0 else verse_section_data["class"][0]
        content = verse_section_data.getText()

        if verse_section_type == "label":
            passage_pointer.update_state(content)
        
        if passage_pointer.state != PassagePointer.IN_PASSAGE:
            return
        
        if paragraph["paragraph"] == False:
            self.add_heading_section(paragraph["doc"], passage_pointer.last_section, passage_pointer) # Add the previous heading
            paragraph["paragraph"] = paragraph["create"](paragraph["doc"])
        
        if verse_section_type == "note":
            content = f"[{footnotes_handler.add_note(verse_section_data)}]"
        
        try:
            paragraph["paragraph"].add_run(content, style=verse_section_type)
        except KeyError:
            print(f"Encountered unexpected verse section type '{verse_section_type}' at v{passage_pointer}. Treating the following as plain text: '{content}'.")

    def add_heading_section(self, paragraph, heading_text, passage_pointer):
        if heading_text and passage_pointer.state == PassagePointer.IN_PASSAGE:
            heading = paragraph.add_paragraph(style="qa")
            heading.add_run(heading_text, style="heading")
        passage_pointer.update_last_section(heading_text)

    def get_copyright_statement(self, yvReader):
        if self.version.custom_copyright_statement:
            copyright_statement = self.version.custom_copyright_statement
        else:
            # Use copyright statement from bible.com directly
            copyright_block = yvReader.find("div", class_=re.compile("ChapterContent_version-copyright"))
            first_section = copyright_block.find("div", recursive=False)
            copyright_statement = first_section.getText()
        return copyright_statement

    def __str__(self):
        end_verse = "end" if self.end_verse == -1 else self.end_verse 
        return f"{self.book_code} {self.chapter}:{self.start_verse}-{end_verse} ({self.version})"

def generate_regular_cafe_handout(book_code, chapter, start_verse=1, end_verse=-1):
    doc = Document()
    add_styles(doc)
    for version in [101, 41, 139, 1819, 73]:
        section_1, section_2 = configure_parallel(doc, page_size="A4", portrait=True)
        Passage(version, book_code, chapter, start_verse, end_verse).write_to(section_1)
        Passage(113, book_code, chapter, start_verse, end_verse).write_to(section_2)
        doc.add_section()

    section = configure_singular(doc, page_size="A5", portrait=True)
    Passage(113, book_code, chapter, start_verse, end_verse).write_to(section)
    
    return doc

def generate_single_page(version, book_code, chapter, start_verse=1, end_verse=-1):
    doc = Document()
    add_styles(doc)
    cell = configure_singular(doc, page_size="A4", portrait=True)
    Passage(version, book_code, chapter, start_verse, end_verse).write_to(cell)
    doc.add_section()

    return doc

def generate_verses_cutout_page(version_names, book_code, chapter, start_verse=1, end_verse=-1):
    doc = Document()
    add_styles(doc)
    for version_name in version_names:
        cell = configure_singular(doc, page_size="A4", portrait=True)
        Passage(resolve_version(version_name), book_code, chapter, start_verse, end_verse).write_to(cell)
        cell.add_paragraph()
    
    return doc

doc = generate_regular_cafe_handout("EXO", 3, 1, 15)
# doc = generate_single_page(73, "EPH", 4)

# doc = generate_verses_cutout_page(["MALDIVIAN"]*2 + ["SIMPLIFIED CHINESE"]*8 + ["ENGLISH"]*10 + ["GERMAN"]*1 + ["JAPANESE"]*1 + ["TRADITIONAL CHINESE"]*1, "JHN", 3, 16, 17)
doc.save("generated/out.docx")